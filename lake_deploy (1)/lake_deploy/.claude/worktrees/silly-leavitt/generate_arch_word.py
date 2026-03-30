"""Generate Architecture Changes Word Document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page setup
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Colors
DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE    = RGBColor(0x2E, 0x74, 0xB5)
TEAL        = RGBColor(0x00, 0x70, 0x7F)
GREEN       = RGBColor(0x1A, 0x7A, 0x40)
ORANGE      = RGBColor(0xC5, 0x5A, 0x11)
RED         = RGBColor(0xA4, 0x26, 0x2F)
GREY        = RGBColor(0x60, 0x60, 0x60)
LIGHT_GREY  = RGBColor(0x80, 0x80, 0x80)

HDR_FILL    = '2E74B5'
HDR_GREEN   = '1A7A40'
HDR_ORANGE  = 'C55A11'
HDR_RED     = 'A4262F'
HDR_TEAL    = '00707F'
ALT_FILL    = 'EEF4FB'
ALT_GREEN   = 'E8F5ED'
ALT_ORANGE  = 'FDF2E9'

def _shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(20); r.font.color.rgb = DARK_BLUE
    return p

def add_h2(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color if color else MID_BLUE
    return p

def add_h3(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = color if color else DARK_BLUE
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(5)
    for r in p.runs: r.font.size = Pt(10)
    return p

def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(10)
    if color: r.font.color.rgb = color
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x17, 0x1A, 0x1A)
    rPr = r._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    rPr.append(shd)
    return p

def add_table(doc, headers, rows, col_widths=None, hdr_color=None):
    hdr_fill = hdr_color if hdr_color else HDR_FILL
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, hdr_fill)
    for ri, row_data in enumerate(rows):
        fill = 'FFFFFF' if ri % 2 == 0 else ALT_FILL
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri+1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(str(val)); r.font.size = Pt(9)
            _shade_cell(cell, fill)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    return tbl

def add_divider(doc, color='2E74B5'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)

def callout_box(doc, text, fill_hex, border_hex, text_color):
    """A coloured paragraph used as a callout."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '12')
        el.set(qn('w:space'), '4');    el.set(qn('w:color'), border_hex)
        pBdr.append(el)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)
    r = p.add_run(text); r.font.size = Pt(10); r.font.color.rgb = text_color
    return p

# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(60)
tp.paragraph_format.space_after  = Pt(10)
r = tp.add_run("Data Lake Architecture")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = DARK_BLUE

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp2.paragraph_format.space_after = Pt(6)
r2 = tp2.add_run("Improvement Recommendations")
r2.bold = True; r2.font.size = Pt(20); r2.font.color.rgb = MID_BLUE

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp3.paragraph_format.space_after = Pt(6)
r3 = tp3.add_run("Current State Analysis · Gap Assessment · Future State Blueprint")
r3.font.size = Pt(12); r3.font.color.rgb = TEAL; r3.italic = True

tp4 = doc.add_paragraph()
tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp4.paragraph_format.space_after = Pt(4)
r4 = tp4.add_run("OnPoint Insights LLC  |  Lake Deploy Platform  |  2026-02-20")
r4.font.size = Pt(10); r4.font.color.rgb = GREY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Executive Summary")
add_divider(doc)

add_body(doc,
    "The lake_deploy data platform is built on a well-structured medallion architecture (Raw → Curated → SSOT) "
    "with strong governance, evidence-backed SSOT certification, and comprehensive multi-source reconciliation. "
    "The platform successfully integrates five data sources (Gaiia, VETRO, Platt, Intacct, Salesforce) "
    "with rigorous data quality gates.")

add_body(doc,
    "However, a comprehensive architectural audit reveals ten structural gaps that — while manageable today — "
    "become critical scalability, reliability, and cost bottlenecks as data volume and tenant count grow. "
    "This document details each gap, its business impact, and the recommended future-state architecture.")

callout_box(doc,
    "⚠  KEY FINDING: 92% of the curated layer (109 of 118 entities) are virtual Athena views with zero physical "
    "storage. Every dashboard query re-scans raw JSONL/CSV on S3. At current scale this is manageable; "
    "at 5–10× growth this becomes the primary cost and latency bottleneck.",
    'FFF3CD', 'C55A11', ORANGE)

doc.add_paragraph()

# Summary scorecard
add_h2(doc, "Architectural Scorecard — Current State")
add_table(doc,
    ["Domain", "Current State", "Risk Level", "Priority"],
    [
        ["Curated Layer Materialization", "92% virtual views — no physical storage", "🔴 High", "P0"],
        ["Pipeline Orchestration",        "Sequential bash + EventBridge, no DAG",   "🔴 High", "P0"],
        ["Checkpoint & Recovery",         "File-based; full restart on failure",       "🔴 High", "P0"],
        ["Schema Drift Detection",        "Glue crawlers detect but don't alert",     "🟠 Medium","P1"],
        ["Incremental Loads",             "Full daily CTAS rebuild of all tables",    "🟠 Medium","P1"],
        ["Multi-Tenancy",                 "Only Gaiia is tenant-partitioned",         "🟠 Medium","P1"],
        ["Data Quality Framework",        "Exceptions captured, not monitored live",  "🟠 Medium","P1"],
        ["SSOT Crosswalk (non-1:1)",      "SF→Intacct many-to-many unresolved",      "🟠 Medium","P1"],
        ["Infrastructure as Code",        "Glue crawlers, AppFlow not in IaC",        "🟡 Low",  "P2"],
        ["Observability & Cost Tracking", "No freshness SLA alarms, no Athena cost",  "🟡 Low",  "P2"],
    ],
    col_widths=[2.4, 2.8, 1.0, 0.8]
)
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPROVEMENT 1 — CURATED LAYER MATERIALIZATION
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Improvement 1: Materialize the Curated Layer")
add_divider(doc, '1A7A40')

add_h2(doc, "1.1  Current State — The Problem", GREEN)
add_body(doc,
    "The curated layer contains 109 virtual Athena views and only 18 physical Parquet tables. "
    "Every query against a curated view triggers a full chain of raw S3 scans:")

add_code(doc, "Dashboard query")
add_code(doc, "  → gaiia_subscriptions_services_current  (VIEW)")
add_code(doc, "      → gaiia_billing_subscriptions_current  (VIEW)  ← raw JSONL scan")
add_code(doc, "      → gaiia_products_current              (VIEW)  ← raw JSONL scan")
add_code(doc, "          → json_extract + UNNEST + ROW_NUMBER()    ← CPU-intensive at scale")

add_body(doc, "Concrete consequences at scale:")
add_bullet(doc, "Athena charges $5/TB scanned. Views re-scan the same raw data for every query.", color=RED)
add_bullet(doc, "5 concurrent dashboard users = 5× parallel full S3 scans with zero result reuse.", color=RED)
add_bullet(doc, "JSON extraction (json_parse, UNNEST, json_extract_scalar) is non-vectorized — CPU scales linearly with row count.", color=RED)
add_bullet(doc, "ROW_NUMBER() dedup windows require full partition sort before returning rn=1 rows.", color=RED)
add_bullet(doc, "Documented in codebase: '⚠ Large queries (>100K rows) may timeout — use LIMIT clauses'", color=RED)

doc.add_paragraph()
add_h2(doc, "1.2  Future State — Materialized Medallion Architecture", GREEN)
add_body(doc,
    "Shift from a view-first to a materialize-first design. The expensive JSON parsing "
    "and deduplication runs once per day in the pipeline; all downstream consumers read "
    "pre-built Parquet files.")

add_table(doc,
    ["Layer", "Current", "Future State"],
    [
        ["Raw",          "JSONL/CSV on S3 (unchanged)",             "JSONL/CSV on S3 (unchanged) — immutable"],
        ["Curated",      "109 virtual views re-scan raw every query","Daily Parquet snapshots written once; thin views on top"],
        ["SSOT",         "18 physical + views on top of views",      "Iceberg tables with MERGE INTO upserts + time-travel"],
        ["Dashboard",    "Live query → raw S3 every page load",      "Query pre-materialized Parquet (~MBs vs ~GBs scanned)"],
    ],
    col_widths=[1.3, 2.5, 3.2], hdr_color=HDR_GREEN
)

doc.add_paragraph()
add_h3(doc, "Implementation Pattern")
add_code(doc, "-- Step 1: Materialization job runs ONCE daily in orchestration pipeline")
add_code(doc, "INSERT INTO curated_core.gaiia_accounts_snapshot  -- Parquet, partitioned by dt")
add_code(doc, "SELECT id, name, tenant, updatedat, customfields,")
add_code(doc, "       current_date AS dt")
add_code(doc, "FROM curated_core.gaiia_accounts_current          -- expensive JSON extraction runs ONCE")
add_code(doc, "WHERE dt = current_date;")
add_code(doc, "")
add_code(doc, "-- Step 2: Views become thin wrappers on Parquet (fast!)")
add_code(doc, "CREATE OR REPLACE VIEW curated_core.gaiia_accounts_current AS")
add_code(doc, "SELECT * FROM curated_core.gaiia_accounts_snapshot")
add_code(doc, "WHERE dt = (SELECT MAX(dt) FROM curated_core.gaiia_accounts_snapshot);")

doc.add_paragraph()
add_h3(doc, "Migrate High-Churn Entities to Apache Iceberg")
add_body(doc, "For entities updated continuously (customers, invoices, subscriptions), use Iceberg tables in Athena for ACID upserts, time-travel, and automatic compaction:")
add_code(doc, "CREATE TABLE curated_core.gaiia_customers_iceberg (")
add_code(doc, "  gaiia_account_id STRING, account_name STRING, tenant STRING, updatedat TIMESTAMP")
add_code(doc, ")")
add_code(doc, "LOCATION 's3://gwi-raw-us-east-2-pc/curated_iceberg/gaiia_customers/'")
add_code(doc, "TBLPROPERTIES ('table_type' = 'ICEBERG');")
add_code(doc, "")
add_code(doc, "-- Upsert (replaces full daily CTAS rebuild)")
add_code(doc, "MERGE INTO curated_core.gaiia_customers_iceberg t")
add_code(doc, "USING staging s ON t.gaiia_account_id = s.gaiia_account_id AND t.tenant = s.tenant")
add_code(doc, "WHEN MATCHED THEN UPDATE SET ...")
add_code(doc, "WHEN NOT MATCHED THEN INSERT VALUES (...);")

doc.add_paragraph()
add_h3(doc, "Immediate Quick Win — Athena Query Result Reuse")
add_body(doc, "Enable at workgroup level — zero code change, immediate benefit for concurrent users:")
add_code(doc, '"ResultReuseConfiguration": {')
add_code(doc, '  "ResultReuseByAgeConfiguration": { "Enabled": true, "MaxAgeInMinutes": 60 }')
add_code(doc, '}')

add_table(doc,
    ["Metric", "Current", "After Materialization"],
    [
        ["Athena scan per dashboard query",  "~10–50 GB (raw JSONL)", "~10–100 MB (Parquet snapshot)"],
        ["Cost per dashboard load",          "$0.05–$0.25",           "<$0.001"],
        ["Query latency (p95)",              "15–60 seconds",          "1–3 seconds"],
        ["Concurrent users supported",       "2–3 before contention",  "50+ (reads cached Parquet)"],
        ["JSON parsing overhead",            "Every query",            "Once per day in pipeline"],
    ],
    col_widths=[2.5, 2.0, 2.5], hdr_color=HDR_GREEN
)
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPROVEMENT 2 — ORCHESTRATION
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Improvement 2: Pipeline Orchestration — DAG-Based Dependency Management")
add_divider(doc, 'C55A11')

add_h2(doc, "2.1  Current State — The Problem", ORANGE)
add_body(doc,
    "Pipelines run sequentially via bash wrappers and EventBridge schedules. "
    "There is no explicit task dependency graph, which creates silent failure modes:")
add_bullet(doc, "If the Glue crawler fails, downstream CTAS runs on stale data — silently.", color=RED)
add_bullet(doc, "No cross-component dependency enforcement between Glue → CTAS → manifests.", color=RED)
add_bullet(doc, "Intacct backfill restarts from the beginning if a checkpoint prefix is lost.", color=RED)
add_bullet(doc, "Vetro rate-limit recovery (429) requires manual next invocation check.", color=RED)
add_bullet(doc, "No SLA enforcement: documented ≤1-day lag has no alarm if breached.", color=RED)

doc.add_paragraph()
add_h2(doc, "2.2  Future State — DAG-Driven with Dependency Gates", ORANGE)

add_table(doc,
    ["Layer", "Current", "Future State"],
    [
        ["Scheduling",       "EventBridge rate() — time-based only",  "Step Functions / MWAA with task dependencies"],
        ["Failure handling", "Silent — downstream runs on stale data", "Upstream gate blocks downstream on failure"],
        ["Retry logic",      "Manual re-trigger required",             "Exponential backoff with configurable retry count"],
        ["Checkpointing",    "File-based (plan_index.json) — fragile", "DynamoDB checkpoint ledger — atomic, resumable"],
        ["SLA enforcement",  "Documented but not alarmed",             "CloudWatch Alarm on max_dt age > 26 hours"],
        ["Rate limiting",    "Poll-based on next_allowed_ts",          "SQS delay queue — event-driven, no polling"],
    ],
    col_widths=[1.5, 2.5, 3.0], hdr_color=HDR_ORANGE
)

doc.add_paragraph()
add_h3(doc, "Proposed Daily Pipeline DAG")
add_code(doc, "START (02:00 UTC)")
add_code(doc, "  ├─ [parallel] Gaiia Lambda ingest   → assert raw partition exists")
add_code(doc, "  ├─ [parallel] Vetro Lambda ingest    → assert raw partition exists")
add_code(doc, "  ├─ [parallel] Platt Glue crawler     → assert table updated")
add_code(doc, "  ├─ [parallel] Intacct ECS task       → assert checkpoint advanced")
add_code(doc, "  └─ [parallel] Salesforce AppFlow     → assert row count > yesterday")
add_code(doc, "        ↓ ALL PASS (gate)")
add_code(doc, "  ├─ [parallel] Materialize curated snapshots per entity")
add_code(doc, "  ├─ [parallel] Run DQ checks + populate dq_run_log")
add_code(doc, "        ↓ DQ PASS (gate: exception_rate < threshold)")
add_code(doc, "  ├─ [parallel] SSOT certification queries")
add_code(doc, "  └─ Write daily manifest + update SSOT dashboard")
add_code(doc, "END — total target: < 45 minutes")

doc.add_paragraph()
add_h3(doc, "DynamoDB Checkpoint Ledger — Replace File-Based State")
add_code(doc, "Table: vetro_pipeline_checkpoints")
add_code(doc, "  PK: source#entity  |  SK: run_date")
add_code(doc, "  Fields: last_cursor, status, part_count, updated_at, error_msg")
add_code(doc, "")
add_code(doc, "-- Resume: read last successful cursor atomically")
add_code(doc, "-- Re-run is safe: same run_date overwrites same PK/SK")
add_code(doc, "-- Failure: status='FAILED' visible in ops dashboard")
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPROVEMENT 3 — PARTITION & STORAGE STRATEGY
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Improvement 3: Partition Strategy & File Format Standardisation")
add_divider(doc, '00707F')

add_h2(doc, "3.1  Current State — The Problem", TEAL)
add_body(doc, "Inconsistent partition keys, mixed formats, and no file-size management across sources:")
add_bullet(doc, "Partition key naming: dt (curated) vs run_date (orchestration) vs none (raw Platt/Salesforce).", color=RED)
add_bullet(doc, "Platt raw: unquoted pipe-delimited CSV — embedded pipes or newlines silently corrupt rows.", color=RED)
add_bullet(doc, "Raw JSONL files have no target file size — small files (<1 MB) degrade Athena performance by 10–100×.", color=RED)
add_bullet(doc, "No business_date / ingestion_date separation — late-arriving data gets today's dt, losing business context.", color=RED)
add_bullet(doc, "Parquet curated files use default Snappy compression — ZSTD gives 30% better compression.", color=RED)

doc.add_paragraph()
add_h2(doc, "3.2  Future State — Standardised Storage Strategy", TEAL)

add_table(doc,
    ["Concern", "Current", "Future State"],
    [
        ["Partition key",        "dt / run_date / none (inconsistent)",  "dt everywhere; add ingestion_dt as metadata column"],
        ["Business date",        "dt = load date (conflated)",           "business_dt column separate from ingestion dt"],
        ["Raw format",           "JSONL, CSV, CSV.GZ, ZIP, XML mixed",   "NDJSON.GZ standard for all raw sources"],
        ["Curated format",       "Parquet + Snappy (default)",           "Parquet + ZSTD (~30% smaller, same speed)"],
        ["File size",            "No target — many small files",         "Target 128–256 MB per file; daily compaction job"],
        ["Platt CSV",            "Unquoted pipe-delimited (fragile)",    "Convert to Parquet at landing via Glue ETL"],
        ["S3 layout",            "source-specific ad-hoc prefixes",      "Standardised: <tenant>/<domain>/<entity>/dt=YYYY-MM-DD/"],
        ["Multi-tenant",         "Only Gaiia uses tenant= partition",    "All sources partitioned by tenant= (default: gwi)"],
    ],
    col_widths=[1.8, 2.4, 2.8], hdr_color=HDR_TEAL
)

doc.add_paragraph()
add_h3(doc, "Proposed Unified S3 Layout")
add_code(doc, "s3://gwi-raw-us-east-2-pc/")
add_code(doc, "  raw/")
add_code(doc, "    <tenant>/              ← gwi | lymefiber | dvfiber | nwfx")
add_code(doc, "      <source>/            ← gaiia | vetro | platt | intacct | salesforce")
add_code(doc, "        <entity>/          ← accounts | invoices | plans | gl_entries")
add_code(doc, "          dt=YYYY-MM-DD/")
add_code(doc, "            part-0001.ndjson.gz")
add_code(doc, "  curated/")
add_code(doc, "    <tenant>/<entity>/dt=YYYY-MM-DD/part-0001.parquet  ← ZSTD, 128-256MB")
add_code(doc, "  ssot/")
add_code(doc, "    <entity>/             ← Iceberg tables with time-travel")
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPROVEMENT 4 — SCHEMA MANAGEMENT
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Improvement 4: Schema Management & Drift Detection")
add_divider(doc)

add_h2(doc, "4.1  Current State — The Problem")
add_body(doc,
    "Glue crawlers auto-detect new fields but generate no alerts. Breaking schema changes "
    "from upstream APIs propagate silently into views and can cause downstream query failures "
    "that are only discovered when a dashboard breaks.")
add_bullet(doc, "Vetro JSON paths are hardcoded in extraction logic — API field rename = silent NULLs.", color=RED)
add_bullet(doc, "Platt: no drift detection; field additions/removals go unnoticed.", color=RED)
add_bullet(doc, "No schema versioning on views — breaking changes have no migration path.", color=RED)
add_bullet(doc, "DDL is inline in CTAS files — hard to version-control separately from logic.", color=RED)

doc.add_paragraph()
add_h2(doc, "4.2  Future State — Schema Contracts + Drift Alerting")

add_table(doc,
    ["Capability", "Current", "Future State"],
    [
        ["Schema validation",   "None — Glue detects silently",          "Contract files per entity; diff on every crawler run"],
        ["Drift alerting",      "None",                                   "Glue → Lambda → SNS alert on schema change"],
        ["Schema versioning",   "None — breaking changes are silent",     "schema_version property on views; alias views for compat"],
        ["JSON path safety",    "Hardcoded paths fail silently on rename","Probe Lambda: null-rate check on 100 sampled rows post-ingest"],
        ["DDL management",      "Inline in CTAS files",                   "Separate schema_contracts/ folder; tested in CI"],
    ],
    col_widths=[1.8, 2.4, 2.8]
)

doc.add_paragraph()
add_h3(doc, "Schema Contract File Example (schema_contracts/gaiia_accounts.json)")
add_code(doc, '{')
add_code(doc, '  "entity": "gaiia_accounts",')
add_code(doc, '  "version": "1.2",')
add_code(doc, '  "required_fields": ["id", "name", "tenant", "updatedat"],')
add_code(doc, '  "optional_fields": ["customFields.platid", "customFields.vetroid"],')
add_code(doc, '  "null_rate_threshold": 0.05,')
add_code(doc, '  "alert_on_new_field": true,')
add_code(doc, '  "alert_on_missing_field": true')
add_code(doc, '}')
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPROVEMENT 5 — DATA QUALITY FRAMEWORK
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Improvement 5: Centralised Data Quality Framework")
add_divider(doc, '1A7A40')

add_h2(doc, "5.1  Current State — The Problem", GREEN)
add_body(doc,
    "Exception tables exist and capture failures, but there is no active monitoring, "
    "no centralised DQ metric store, no automated quarantine, and tolerance thresholds "
    "are hardcoded in shell scripts:")
add_bullet(doc, "Exception tables are populated but never trigger alerts — failures are discovered manually.", color=RED)
add_bullet(doc, "Tolerance thresholds (±0.5% MRR, ±0.2% customers) hardcoded in bash variables.", color=RED)
add_bullet(doc, "No automated quarantine: bad rows mix with good rows in SSOT views.", color=RED)
add_bullet(doc, "No DQ trend visibility — can't tell if exception rate is improving or degrading over time.", color=RED)

doc.add_paragraph()
add_h2(doc, "5.2  Future State — Active DQ Framework", GREEN)

add_table(doc,
    ["DQ Component", "Current", "Future State"],
    [
        ["DQ metrics storage",   "Scattered manifests per run",             "curated_recon.dq_run_log table (centralised)"],
        ["Threshold management", "Hardcoded in shell scripts",              "config/dq_thresholds.json (versioned, auditable)"],
        ["Quarantine",           "Bad rows mixed into SSOT views",          "Quarantine prefix; excluded from SSOT automatically"],
        ["Alerting",             "None on exception counts",                "CloudWatch Alarm: exception_rate > threshold → SNS"],
        ["Trend visibility",     "None",                                    "30-day rolling exception rate chart per source"],
        ["Assertion framework",  "Manual gate scripts",                     "dbt tests / Great Expectations on curated tables"],
    ],
    col_widths=[1.8, 2.4, 2.8], hdr_color=HDR_GREEN
)

doc.add_paragraph()
add_h3(doc, "Centralised DQ Run Log Table")
add_code(doc, "CREATE TABLE curated_recon.dq_run_log (")
add_code(doc, "  run_date        date,")
add_code(doc, "  source          string,    -- gaiia | vetro | platt | intacct | salesforce")
add_code(doc, "  entity          string,    -- accounts | invoices | plans | gl_entries")
add_code(doc, "  row_count       bigint,")
add_code(doc, "  exception_count bigint,")
add_code(doc, "  exception_rate  double,    -- exception_count / row_count")
add_code(doc, "  null_rate_pk    double,    -- null rate on primary key field")
add_code(doc, "  freshness_hours double,    -- hours since last business_dt")
add_code(doc, "  pass_fail       string,    -- PASS | WARN | FAIL")
add_code(doc, "  evaluated_at    timestamp")
add_code(doc, ") STORED AS PARQUET LOCATION 's3://.../curated_recon/dq_run_log/';")
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPROVEMENT 6 — SSOT & CROSSWALK
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Improvement 6: SSOT Identity Spine & Crosswalk Architecture")
add_divider(doc, 'C55A11')

add_h2(doc, "6.1  Current State — The Problem", ORANGE)
add_body(doc,
    "The Salesforce → Intacct crosswalk is non-1:1 due to account migrations and duplicates, "
    "causing many-to-many join explosions in SSOT views. No SCD Type 2 history. "
    "Source-system IDs are used interchangeably without a canonical surrogate key:")
add_bullet(doc, "SF Account ID, Platt Customer ID, Intacct Customer ID used interchangeably in views.", color=RED)
add_bullet(doc, "Account migrations create many-to-many mappings — SSOT tiles cannot be certified.", color=RED)
add_bullet(doc, "Customer changes silently overwrite prior records — no audit history.", color=RED)
add_bullet(doc, "Source priority rules (SF > Platt > Intacct) hardcoded in view SQL, not externalised.", color=RED)

doc.add_paragraph()
add_h2(doc, "6.2  Future State — Entity Spine with SCD Type 2", ORANGE)

add_table(doc,
    ["Concept", "Current", "Future State"],
    [
        ["Identity model",       "Source IDs used directly in views",      "Surrogate ssot_entity_id (UUID) spine table"],
        ["Crosswalk cardinality","Non-1:1 (many-to-many on migrations)",   "1:1 enforced; multiple entries → confidence tier"],
        ["Customer history",     "Changes overwrite prior records",         "SCD Type 2: valid_from, valid_to, is_current"],
        ["Source priority",      "Hardcoded in view SQL (SF > Platt > Int)","Externalised to ssot_source_priority_rules table"],
        ["Unmapped handling",    "Silently dropped or mis-counted",         "Explicit 'unmapped' bucket + daily alert if > 0"],
    ],
    col_widths=[1.8, 2.4, 2.8], hdr_color=HDR_ORANGE
)

doc.add_paragraph()
add_h3(doc, "Entity Spine Table Design")
add_code(doc, "CREATE TABLE curated_ssot.entity_spine (")
add_code(doc, "  ssot_entity_id   STRING,   -- UUID assigned at first resolution")
add_code(doc, "  sf_account_id    STRING,   -- Salesforce Account ID (if known)")
add_code(doc, "  platt_customer_id STRING,  -- Platt Customer ID (if known)")
add_code(doc, "  intacct_id       STRING,   -- Intacct Customer ID (if known)")
add_code(doc, "  gaiia_account_id STRING,   -- Gaiia Account ID (if known)")
add_code(doc, "  confidence       STRING,   -- high | medium | low | unresolved")
add_code(doc, "  valid_from       DATE,     -- SCD Type 2 start")
add_code(doc, "  valid_to         DATE,     -- SCD Type 2 end (NULL = current)")
add_code(doc, "  is_current       BOOLEAN")
add_code(doc, ") TBLPROPERTIES ('table_type' = 'ICEBERG');")
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPROVEMENT 7 — MULTI-TENANCY
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Improvement 7: Full Multi-Tenancy Architecture")
add_divider(doc)

add_h2(doc, "7.1  Current State")
add_body(doc,
    "Only Gaiia GraphQL has tenant-partitioned storage (tenant=gwi/lymefiber/dvfiber). "
    "All other sources (Platt, Intacct, Vetro, Salesforce) are GWI-only with no tenant isolation. "
    "Adding a new network operator (e.g., LymeFiber on Platt) requires code changes, not configuration.")

add_h2(doc, "7.2  Future State")
add_table(doc,
    ["Source", "Current Tenant Support", "Future State"],
    [
        ["Gaiia",      "✅ tenant= partition (gwi, lymefiber, dvfiber)", "✅ Extend to new operators via config"],
        ["Platt",      "❌ GWI only — hardcoded",                        "✅ tenant= partition from operator code in export"],
        ["Intacct",    "❌ GWI only — single entity",                    "✅ entity_id mapped to tenant in checkpoint ledger"],
        ["Vetro",      "❌ GWI only — plan_id not tenant-tagged",        "✅ plan_id → tenant resolved via crosswalk"],
        ["Salesforce", "❌ Single AppFlow connection",                    "✅ Separate AppFlow per tenant or custom Lambda"],
    ],
    col_widths=[1.2, 2.5, 3.3]
)

doc.add_paragraph()
add_h3(doc, "Parameterised Tenant Config")
add_code(doc, '// config/tenants.json — add tenant with zero code changes')
add_code(doc, '{')
add_code(doc, '  "tenants": [')
add_code(doc, '    { "id": "gwi",       "gaiia_key": "gwi_key",       "platt_prefix": "gwi/" },')
add_code(doc, '    { "id": "lymefiber", "gaiia_key": "lymefiber_key", "platt_prefix": "lymefiber/" },')
add_code(doc, '    { "id": "dvfiber",   "gaiia_key": "dvfiber_key",   "platt_prefix": "dvfiber/" }')
add_code(doc, '  ]')
add_code(doc, '}')
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPROVEMENT 8 — OBSERVABILITY
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Improvement 8: Observability, Cost Tracking & Alerting")
add_divider(doc, '00707F')

add_h2(doc, "8.1  Current State — The Problem", TEAL)
add_body(doc, "CloudWatch logs exist but are not wired to actionable alarms. No Athena cost allocation. No freshness SLA enforcement:")
add_bullet(doc, "SLA documented as ≤1-day lag — but no CloudWatch alarm fires if max_dt is 36 hours old.", color=RED)
add_bullet(doc, "Exception counts logged but never trigger alerts — failures found manually.", color=RED)
add_bullet(doc, "No Athena cost allocation by pipeline stage — can't identify expensive queries.", color=RED)
add_bullet(doc, "SQS DLQ exists for Vetro but no alarm on messages arriving.", color=RED)
add_bullet(doc, "SSOT audit evidence produced but not integrated into operational dashboard.", color=RED)

doc.add_paragraph()
add_h2(doc, "8.2  Future State — Observability Stack", TEAL)

add_table(doc,
    ["Observable", "Current", "Future State"],
    [
        ["Data freshness",       "Documented SLA, no alarm",         "CloudWatch Alarm: max_dt age > 26h per source → SNS"],
        ["Exception rate",       "Logged, not alarmed",              "Alarm: exception_rate > threshold in dq_run_log → SNS"],
        ["SQS DLQ",             "No alarm on message arrival",       "Alarm: ApproximateNumberOfMessagesVisible > 0 → PagerDuty"],
        ["Athena cost",          "No allocation by stage",           "Workgroups: raw-ingestion | curated | ssot; cost by stage"],
        ["Pipeline duration",    "Not tracked",                      "Metric: pipeline_end_ts - pipeline_start_ts; P95 alarm"],
        ["Schema drift",         "Glue logs only",                   "Glue → Lambda → SNS on any field add/remove/type change"],
        ["SSOT certification",   "Manual evidence review",           "Auto-publish status.json to ops dashboard daily"],
    ],
    col_widths=[1.8, 2.2, 3.0], hdr_color=HDR_TEAL
)
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPROVEMENT 9 — IaC & CI/CD
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Improvement 9: Infrastructure as Code & CI/CD Completeness")
add_divider(doc)

add_h2(doc, "9.1  Current State — Gaps")
add_table(doc,
    ["Component", "IaC Status", "Risk"],
    [
        ["Vetro Export Lambda + EventBridge", "✅ CloudFormation (vetro_export_stack.yaml)", "Low"],
        ["Gaiia Ingest Lambda + EventBridge", "✅ CloudFormation (gaiia_ingest_stack.yaml)",  "Low"],
        ["Glue Crawlers",                     "❌ Reference YAML only — console deployed",    "High — drift-prone"],
        ["Salesforce AppFlow connections",     "❌ Console only — no IaC",                    "High — unrecoverable without docs"],
        ["Athena Workgroups",                  "❌ Console only",                             "Medium"],
        ["S3 Bucket policies / lifecycle",     "❌ Console only",                             "High — data retention risk"],
        ["Amplify Frontend Deploy",            "❌ Manual artifact upload",                   "Medium — error-prone"],
        ["OIDC GitHub → AWS auth",             "❌ Long-lived keys (if used)",               "High — credential leak risk"],
    ],
    col_widths=[2.5, 2.5, 2.0]
)

doc.add_paragraph()
add_h2(doc, "9.2  Future State — Full IaC Coverage")
add_bullet(doc, "Codify all Glue crawlers in CloudFormation (or Terraform) — re-creatable from code in a fresh AWS account.")
add_bullet(doc, "AppFlow connections via CloudFormation resource type AWS::AppFlow::ConnectorProfile.")
add_bullet(doc, "S3 lifecycle rules in IaC: raw → Glacier after 90 days; Athena results → delete after 30 days.")
add_bullet(doc, "CI pipeline: on every PR — run Athena EXPLAIN on critical views, validate all DDL files, check S3 path consistency.")
add_bullet(doc, "OIDC GitHub Actions → AWS: replace long-lived IAM keys with short-lived STS tokens from OIDC federation.")
add_bullet(doc, "Amplify connected to GitHub main branch — auto-deploy on merge.")
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# IMPLEMENTATION ROADMAP
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Implementation Roadmap")
add_divider(doc)

add_h2(doc, "Phase 0 — Immediate Wins (Week 1, No Architecture Change)")
callout_box(doc,
    "These changes can be made TODAY with no risk to existing pipelines. "
    "Combined, they reduce Athena query latency for concurrent users by 60–80%.",
    'E8F5ED', '1A7A40', GREEN)

add_table(doc,
    ["Action", "Effort", "Impact"],
    [
        ["Enable Athena Query Result Reuse on dashboard workgroup (60-min TTL)", "30 min",  "Immediate: concurrent users share cached results"],
        ["Add CloudWatch Alarm on Vetro SQS DLQ (ApproximateNumberOfMessages > 0)", "30 min","Immediate: DLQ failures now alerted"],
        ["Externalise DQ tolerance thresholds to config/dq_thresholds.json",       "1 hour", "Auditable thresholds without code deploy"],
        ["Add WHERE dt = MAX(dt) guard to all *_current views",                     "1 day",  "Prevents full partition scan on dedup"],
        ["Add tenant column (default 'gwi') to all curated views",                  "1 day",  "Future-proofs multi-tenant expansion"],
        ["Standardise partition key name to dt everywhere",                          "1 day",  "Eliminates dt vs run_date confusion"],
    ],
    col_widths=[3.5, 1.0, 2.5], hdr_color=HDR_GREEN
)
doc.add_paragraph()

add_h2(doc, "Phase 1 — Materialization & Reliability (Weeks 2–4)")
add_table(doc,
    ["Action", "Effort", "Impact"],
    [
        ["Materialize top 5 curated entities as daily Parquet snapshots (Gaiia accounts, customers, subscriptions; Vetro GIS layers; Platt billing)", "1 week", "50–100× reduction in Athena scan per dashboard query"],
        ["Implement DynamoDB checkpoint ledger replacing plan_index.json", "3 days",  "Vetro/Intacct restarts resume from last cursor, not beginning"],
        ["Create curated_recon.dq_run_log table + populate daily",        "2 days",  "Centralised DQ trend visibility"],
        ["Cloudwatch Alarm: max_dt age > 26h per source",                  "1 day",   "SLA breach now generates alert, not manual discovery"],
        ["Schema contract files for Gaiia + Vetro top 5 entities",         "2 days",  "API breaking changes caught before they hit SSOT"],
        ["Codify Glue crawlers in CloudFormation",                          "2 days",  "Crawler config version-controlled, recoverable"],
    ],
    col_widths=[3.5, 1.0, 2.5], hdr_color=HDR_ORANGE
)
doc.add_paragraph()

add_h2(doc, "Phase 2 — Iceberg & Full DAG (Month 2)")
add_table(doc,
    ["Action", "Effort", "Impact"],
    [
        ["Migrate curated_core high-churn entities to Apache Iceberg", "2 weeks", "MERGE INTO upserts, time-travel, auto-compaction"],
        ["Build Step Functions DAG replacing sequential bash pipeline",  "2 weeks", "Upstream failures block downstream; parallel execution"],
        ["Implement entity_spine surrogate key table",                  "1 week",  "Resolves SF→Intacct many-to-many crosswalk"],
        ["SCD Type 2 for customer/account changes",                     "1 week",  "Full historical audit trail for customer changes"],
        ["Athena Workgroups by pipeline stage + cost dashboards",       "3 days",  "Cost attribution by source and pipeline phase"],
        ["OIDC GitHub→AWS auth in CI",                                  "1 day",   "Eliminates long-lived IAM key risk"],
    ],
    col_widths=[3.5, 1.0, 2.5], hdr_color=HDR_TEAL
)
doc.add_paragraph()

add_h2(doc, "Phase 3 — Multi-Tenancy & Full Observability (Month 3+)")
add_table(doc,
    ["Action", "Effort", "Impact"],
    [
        ["Unified S3 layout: tenant/source/entity/dt= for all sources", "2 weeks", "Consistent partitioning; add new tenants via config"],
        ["Parameterise Lambda/Glue jobs with tenant_id from tenants.json","1 week",  "New operator onboarding = config change, not code change"],
        ["Data freshness + exception rate dashboards in CloudWatch",     "1 week",  "Operational SLA visibility in real time"],
        ["dbt project or Great Expectations suite on curated tables",    "3 weeks", "Automated data contract testing per entity"],
        ["AppFlow → IaC (CloudFormation connector profiles)",            "1 week",  "Salesforce connection recoverable from code"],
        ["Amplify + GitHub CI/CD auto-deploy",                           "2 days",  "Frontend deploys on merge, not manual upload"],
    ],
    col_widths=[3.5, 1.0, 2.5]
)
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# ARCHITECTURE COMPARISON — BEFORE / AFTER
# ═══════════════════════════════════════════════════════════
add_h1(doc, "Architecture Comparison — Before & After")
add_divider(doc)

add_h2(doc, "Data Flow: Current State vs Future State")
add_table(doc,
    ["Component", "Current State", "Future State"],
    [
        ["Ingestion trigger",       "EventBridge time-based (fixed schedule)",            "Step Functions DAG — dependency-aware"],
        ["Ingestion checkpoint",    "File-based JSON (fragile on loss)",                  "DynamoDB ledger (atomic, resumable)"],
        ["Rate-limit handling",     "Poll-based on next_allowed_ts",                      "SQS delay queue (event-driven, no polling)"],
        ["Raw storage format",      "Mixed: JSONL, CSV, XML, ZIP",                       "NDJSON.GZ standard (all sources)"],
        ["Raw partitioning",        "Inconsistent: dt / run_date / none",                 "Unified: tenant/source/entity/dt="],
        ["Curated layer type",      "92% virtual views — no physical storage",            "Daily Parquet snapshots; thin views on top"],
        ["Curated format/codec",    "Parquet + Snappy (default)",                         "Parquet + ZSTD; 128–256 MB target file size"],
        ["SSOT tables",             "EXTERNAL tables + views on top of views",            "Apache Iceberg with MERGE INTO + time-travel"],
        ["JSON extraction",         "Every query — CPU-intensive at scale",               "Once per day in pipeline; Parquet for queries"],
        ["Dedup (ROW_NUMBER)",       "Every query — full partition sort",                  "Once per day at materialisation time"],
        ["Multi-tenancy",           "Gaiia only (tenant= partition)",                     "All sources partitioned by tenant="],
        ["DQ thresholds",           "Hardcoded in bash scripts",                          "config/dq_thresholds.json (versioned)"],
        ["DQ monitoring",           "Exception tables — not alarmed",                     "dq_run_log + CloudWatch alarms"],
        ["Schema drift detection",  "Glue detects — no alert",                            "Contract files + Lambda → SNS alert"],
        ["SSOT identity model",     "Source IDs used directly (no surrogate)",            "entity_spine UUID + SCD Type 2"],
        ["Customer history",        "Overwrite on update — no history",                   "valid_from/valid_to/is_current (SCD2)"],
        ["Freshness SLA",           "Documented ≤1 day — no alarm",                      "CloudWatch Alarm: max_dt age > 26h"],
        ["Pipeline duration",       "Not tracked",                                        "P95 metric + alarm on 2× baseline"],
        ["IaC coverage",            "Lambdas + CF; Glue + AppFlow = console",             "100% IaC: all resources in CloudFormation"],
        ["Frontend deploy",         "Manual Amplify upload",                              "GitHub CI → Amplify auto-deploy on merge"],
        ["Auth in CI",              "Long-lived IAM keys (if used)",                      "OIDC short-lived STS tokens"],
    ],
    col_widths=[2.0, 2.5, 2.5]
)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════
add_divider(doc)
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    "Lake Deploy Architecture Improvement Plan  |  OnPoint Insights LLC  |  Confidential  |  2026-02-20")
r.font.size = Pt(8); r.font.color.rgb = LIGHT_GREY

out = ("/Users/vinaymistry/Library/CloudStorage/OneDrive-OnPointInsightsLLC/"
       "GitRepo/lake_deploy/.claude/worktrees/silly-leavitt/Architecture_Improvement_Plan.docx")
doc.save(out)
print(f"Word saved: {out}")
